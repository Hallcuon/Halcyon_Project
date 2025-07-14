from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import datetime

def create_report():
    doc = Document()
    
    # Налаштування полів сторінки
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.4)

    # Титульна сторінка
    title = doc.add_paragraph()
    title_run = title.add_run('ЗВІТ\nпро проходження виробничої практики')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Тема
    theme = doc.add_paragraph()
    theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    theme_run = theme.add_run('\nТема: "Розробка інтерактивної RPG-гри на основі веб-карти з використанням React та Django"')
    theme_run.font.size = Pt(14)

    # Інформація про місце практики
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info.add_run('\nМісце практики: Проект "Halcyon\'s Project"\n')
    info.add_run('Термін практики: 01.07.2025 - 14.07.2025\n\n')

    # Інформація про виконавця та керівників
    doc.add_paragraph('Виконав:\nСтудент групи [Номер групи]\n[ПІБ студента]\n\n')
    doc.add_paragraph('Керівник практики від університету:\n[Посада, науковий ступінь]\n[ПІБ керівника]\n\n')
    doc.add_paragraph('Керівник практики від підприємства:\n[Посада]\n[ПІБ керівника]\n\n')

    # Місто та рік
    bottom = doc.add_paragraph()
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom.add_run('Київ - 2025')

    # Розрив сторінки перед змістом
    doc.add_page_break()

    # Зміст
    doc.add_heading('ЗМІСТ', level=1)
    
    # Зберігаємо документ
    doc.save('d:/MapProject/report.docx')

if __name__ == "__main__":
    create_report()